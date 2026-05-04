from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = "teacher-google-sheet-guide.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header_fill="EAF3EE"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
                    run.font.size = Pt(10)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        if level == 1:
            run.font.color.rgb = RGBColor(20, 83, 45)
    return heading


def add_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r2 = p.add_run(text[len(bold_prefix):])
        runs = [r1, r2]
    else:
        runs = [p.add_run(text)]
    for run in runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(10.5)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(60, 70, 80)


def add_key_table(doc):
    add_heading(doc, "ภาพรวมการลิงก์ข้อมูล", 1)
    add_paragraph(
        doc,
        "Google Sheet นี้ทำหน้าที่เหมือนฐานข้อมูลของระบบ หน้าผู้เรียนจะอ่านข้อมูลจากหลาย tab แล้วนำมาประกอบกันด้วย ID ดังนั้น ID ต้องสะกดตรงกันทุกตัว",
    )
    rows = [
        ["Tab", "ID หลัก", "ใช้เชื่อมกับ", "หน้าที่"],
        ["LineProfiles", "lineProfileId", "Users.lineProfileId", "เก็บข้อมูลจาก LINE อัตโนมัติหลังผู้เรียน login"],
        ["Users", "userId", "Enrollments.userId", "รายชื่อผู้เรียน/ผู้ใช้หลักของระบบ"],
        ["Courses", "courseId", "Enrollments.courseId", "ข้อมูลคอร์ส เช่น รายครั้งหรือรายชั่วโมง"],
        ["Enrollments", "enrollmentId", "Lessons / Attendances", "การซื้อคอร์สของผู้เรียน 1 คน ต่อ 1 คอร์ส"],
        ["Lessons", "lessonId", "enrollmentId", "ตารางเรียนที่จะเกิดขึ้น"],
        ["Attendances", "attendanceId", "enrollmentId", "ประวัติการเข้าเรียนและจำนวนที่ใช้ไป"],
    ]
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            set_cell_text(table.cell(i, j), value, bold=(i == 0))
    style_table(table)


def add_id_rules(doc):
    add_heading(doc, "กติกาการตั้ง ID", 1)
    add_bullets(
        doc,
        [
            "ใช้ตัวอักษรอังกฤษ ตัวเลข และขีดกลางเท่านั้น เช่น user-001 ห้ามมีเว้นวรรค",
            "ID ต้องไม่ซ้ำภายใน tab เดียวกัน",
            "เมื่อ ID ถูกนำไปใช้แล้ว ห้ามเปลี่ยนย้อนหลัง เพราะข้อมูล tab อื่นอ้างถึง ID นั้นอยู่",
            "แนะนำให้ copy ID จาก tab ต้นทางไปวางใน tab ปลายทาง เพื่อลดการพิมพ์ผิด",
            "ถ้าไม่แน่ใจ ให้เพิ่มเลขลำดับต่อท้าย เช่น user-004, course-007, enr-012",
        ],
    )

    rows = [
        ["ชนิดข้อมูล", "รูปแบบแนะนำ", "ตัวอย่าง"],
        ["ผู้เรียน", "user-เลขลำดับ", "user-001"],
        ["คอร์ส", "course-เลขลำดับ", "course-001"],
        ["การสมัคร/ซื้อคอร์ส", "enr-เลขลำดับ", "enr-001"],
        ["ตารางเรียน", "les-เลขลำดับ", "les-001"],
        ["ประวัติการเข้าเรียน", "att-เลขลำดับ", "att-001"],
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            set_cell_text(table.cell(i, j), value, bold=(i == 0))
    style_table(table, "EEF2F5")


def add_workflows(doc):
    add_heading(doc, "ขั้นตอนเพิ่มข้อมูลที่ครูใช้บ่อย", 1)

    add_heading(doc, "1. เพิ่มผู้เรียนใหม่", 2)
    add_bullets(
        doc,
        [
            "ให้ผู้เรียนเปิด LIFF/login LINE ก่อน 1 ครั้ง เพื่อให้ระบบสร้างแถวใน tab LineProfiles",
            "ไปที่ tab LineProfiles แล้ว copy ค่า lineProfileId ของผู้เรียน",
            "ไปที่ tab Users แล้วเพิ่มแถวใหม่: userId, lineProfileId, displayName, pictureUrl, role",
            "role สำหรับผู้เรียนให้ใส่ STUDENT",
        ],
    )
    add_code(doc, "Users: user-001 | <lineProfileId จาก LineProfiles> | คุณซัน | | STUDENT")

    add_heading(doc, "2. เพิ่มคอร์สใหม่", 2)
    add_bullets(
        doc,
        [
            "ไปที่ tab Courses",
            "เพิ่ม courseId ใหม่ เช่น course-001",
            "courseType ใส่ CLASS ถ้านับเป็นครั้ง หรือ HOUR ถ้านับเป็นชั่วโมง",
            "totalClasses คือจำนวนทั้งหมดของคอร์สนั้น เช่น 10",
        ],
    )
    add_code(doc, "Courses: course-001 | Private Course 10 Classes | CLASS | 10")

    add_heading(doc, "3. ให้ผู้เรียนซื้อ/สมัครคอร์ส", 2)
    add_bullets(
        doc,
        [
            "ไปที่ tab Enrollments",
            "สร้าง enrollmentId ใหม่ เช่น enr-001",
            "ใส่ userId จาก tab Users",
            "ใส่ courseId จาก tab Courses",
            "purchasedClasses คือจำนวนที่ซื้อทั้งหมด",
            "remainingClasses คือจำนวนคงเหลือปัจจุบัน",
            "status ใช้ ACTIVE, PAUSED, COMPLETED หรือ CANCELLED",
        ],
    )
    add_code(doc, "Enrollments: enr-001 | user-001 | course-001 | 10 | 10 | ACTIVE")

    add_heading(doc, "4. บันทึกประวัติหลังเรียนจบ", 2)
    add_bullets(
        doc,
        [
            "ไปที่ tab Attendances",
            "สร้าง attendanceId ใหม่ เช่น att-001",
            "ใส่ enrollmentId ของคอร์สที่ผู้เรียนใช้เรียน",
            "checkedInAt ใส่วันเวลา เช่น 2026-05-04T10:05:00+07:00",
            "classesUsed ใส่จำนวนที่ใช้ เช่น 1",
            "กลับไปลด remainingClasses ใน tab Enrollments ให้ถูกต้อง เช่น จาก 10 เหลือ 9",
        ],
    )
    add_code(doc, "Attendances: att-001 | enr-001 | Teacher A | 2026-05-04T10:05:00+07:00 | 1 | เรียนครั้งที่ 1")


def add_examples(doc):
    add_heading(doc, "ตัวอย่างข้อมูลที่ลิงก์กันครบ 1 ชุด", 1)
    rows = [
        ["Tab", "ตัวอย่างข้อมูล"],
        ["Users", "user-001 | line-profile-abc | คุณซัน | | STUDENT"],
        ["Courses", "course-001 | Private Course 10 Classes | CLASS | 10"],
        ["Enrollments", "enr-001 | user-001 | course-001 | 10 | 4 | ACTIVE"],
        ["Attendances", "att-001 | enr-001 | Teacher A | 2026-05-04T10:05:00+07:00 | 1 | เรียนครั้งที่ 1"],
        ["Attendances", "att-002 | enr-001 | Teacher A | 2026-05-06T10:05:00+07:00 | 1 | เรียนครั้งที่ 2"],
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            set_cell_text(table.cell(i, j), value, bold=(i == 0))
    style_table(table)
    add_paragraph(
        doc,
        "จากตัวอย่างนี้หน้าเว็บจะรู้ว่า user-001 เรียน course-001 ผ่าน enr-001 และมีประวัติการเรียนจาก Attendances ที่อ้างถึง enr-001",
    )


def add_checklist(doc):
    add_heading(doc, "Checklist ก่อนบอกว่าข้อมูลพร้อมใช้", 1)
    add_bullets(
        doc,
        [
            "Users.userId มีค่า และไม่ซ้ำ",
            "Enrollments.userId ตรงกับ Users.userId",
            "Enrollments.courseId ตรงกับ Courses.courseId",
            "Lessons.enrollmentId และ Attendances.enrollmentId ตรงกับ Enrollments.enrollmentId",
            "courseType ใส่เป็น CLASS หรือ HOUR เท่านั้น",
            "status ใส่เป็น ACTIVE, PAUSED, COMPLETED หรือ CANCELLED เท่านั้น",
            "remainingClasses อัปเดตหลังเพิ่ม Attendances ทุกครั้ง",
        ],
    )


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("คู่มือดูแล Google Sheet สำหรับระบบคอร์สเรียน")
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(20, 83, 45)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("สำหรับคุณครู/ผู้ดูแลที่ต้องเพิ่มนักเรียน คอร์ส ตารางเรียน และประวัติการเข้าเรียน")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(90, 105, 120)

    add_paragraph(
        doc,
        "หลักคิดสำคัญ: ระบบใช้ userId เป็นตัวหลักของผู้เรียน ส่วน lineProfileId มีไว้เชื่อมบัญชี LINE กับผู้เรียนใน tab Users เท่านั้น",
        bold_prefix="หลักคิดสำคัญ:",
    )

    add_key_table(doc)
    add_id_rules(doc)
    add_workflows(doc)
    add_examples(doc)
    add_checklist(doc)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
